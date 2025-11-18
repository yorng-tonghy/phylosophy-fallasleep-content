from docx import Document
import time
import openai
import math

# Configuration
openai.api_base = "https://openrouter.ai/api/v1"
openai.api_key = "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"
openai.headers = {
    "HTTP-Referer": "http://localhost",
    "X-Title": "DeepSeek Content Generator"
}

def generate_topics(main_topic):
    """Generate 100 numbered subtopics with the complete original prompt"""
    prompt = f"""Please generate exactly 100 topics about {main_topic}. 
Format them as a numbered list (1. Topic 1, 2. Topic 2, etc.). 
Make sure there are exactly 100 items, covering all key aspects of {main_topic}."""
    
    print(f"Generating 100 topics about '{main_topic}'...")
    try:
        response = openai.ChatCompletion.create(
            model="deepseek/deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2500
        )
        topics = []
        for line in response.choices[0].message.content.split('\n'):
            if line.strip():
                parts = line.split('. ', 1)
                if len(parts) > 1:
                    topics.append(parts[1].strip())
        return topics[:100]  # still limited to 20
    except Exception as e:
        print(f"Error generating topics: {str(e)}")
        return []

def generate_paragraph(topic):
    """Generate content with the complete original soothing narration prompt"""
    prompt = f"""You are a calm, thoughtful narrator. For the topic "{topic}", write a single, flowing paragraph of about 200 words that:
1. Begins with "You" not really begin just random use "You"
2. Uses natural, conversational language, 
3. Has soft pacing and gentle clarity
4. Provides meaningful insight without jargon
5. Feels like a personal conversation with a curious friend
6. Includes metaphors or simple visual language where appropriate
7. Maintains a soothing, reassuring tone throughout
8. ask quesion or give example if possible then answer it

Write as if explaining something meaningful to someone during a quiet moment of reflection."""

    try:
        response = openai.ChatCompletion.create(
            model="deepseek/deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Error generating content: {str(e)}")
        return f"[Content generation failed for: {topic}]"

def create_document(topics, paragraphs, filename="Thoughtful_Reflections.docx"):
    """Create a single Word document with all topics and paragraphs"""
    doc = Document()
    doc.add_heading("Thoughtful Reflections", level=0)

    for i, (topic, content) in enumerate(zip(topics, paragraphs), 1):
        doc.add_heading(f"{i}. {topic}", level=1)
        doc.add_paragraph(content)
        doc.add_paragraph("")  # spacing

    doc.save(filename)
    print(f"Saved: {filename}")

def main():
    print("DeepSeek Content Generator".center(50, '='))
    main_topic = input("\nEnter your main topic (e.g., life, emotions, human nature): ").strip()
    
    if not main_topic:
        print("Please enter a valid topic.")
        return
    
    print("\nGenerating 100 subtopics...")
    topics = generate_topics(main_topic)
    
    if not topics:
        print("Failed to generate topics. Please check your API key and connection.")
        return
    
    print(f"\nNow generating paragraphs for {len(topics)} topics...")

    paragraphs = []
    for i, topic in enumerate(topics, 1):
        print(f"Progress: {i}/{len(topics)} - {topic[:50]}...")
        paragraphs.append(generate_paragraph(topic))
        time.sleep(1.2)  # Respect rate limits
    
    create_document(topics, paragraphs)

if __name__ == "__main__":
    main()
